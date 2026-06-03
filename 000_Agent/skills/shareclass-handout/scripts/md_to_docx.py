import os
import re
import copy
import argparse
import docx
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

def apply_font(run, font_name="芫荽", font_size_pt=11, bold=False, color_rgb=None):
    """
    強制在 Word XML 底層設定 Runs 的字型為指定字型
    """
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def add_formatted_text(p, text, default_bold=False, font_name="芫荽", font_size_pt=11, color_rgb=None):
    """
    解析 Markdown 粗體 ** 語法，並依序將 Runs 加入段落中，強制套用指定字型。
    """
    if not text:
        return
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner_text = part[2:-2]
            run = p.add_run(inner_text)
            apply_font(run, font_name=font_name, font_size_pt=font_size_pt, bold=True, color_rgb=color_rgb)
        else:
            run = p.add_run(part)
            apply_font(run, font_name=font_name, font_size_pt=font_size_pt, bold=default_bold, color_rgb=color_rgb)

def format_paragraph(p, space_before_pt=0, space_after_pt=0, line_spacing=1.0, keep_with_next=False):
    """
    設定段落格式，預設段落後間距改為 0，且行距預設為 1.0 (單行行距)
    """
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next

def clean_document_content(doc):
    """
    精確清空範本內容，保留 w:sectPr（版面、邊界、頁首頁尾），並維持 Word XML 規範順序
    """
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    
    # 清除 body 內所有不是 sectPr 的子元素
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)
            
    # 依據 Word 規格，在 sectPr 之前必須有一個 w:p 作為基礎段落
    p = OxmlElement('w:p')
    if sectPr is not None:
        sectPr.addprevious(p)
    else:
        body.append(p)
        
    # 重設新文件的段落列表快取
    doc._paragraphs = []
    
def append_element_safely(doc, element):
    """
    精確將段落或表格元素插入到 w:sectPr 之前，防止破壞 Word 頁面配置結構
    """
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(element)
    else:
        body.append(element)

def extract_stamps(doc):
    """
    從原版 3W 講義中提取代表性元數據的樣式印章 (Stamps)
    """
    stamps = {
        'h1': None,
        'h2': None,
        'h3': None,
        'q': None,
        'normal': None,
        'blank': None,
        'merge_table': None,
        'md_table': None,
        'quiz_table': None
    }
    
    for p in doc.paragraphs:
        text_s = p.text.strip()
        if not stamps['h1'] and p.style.name.startswith('Heading 1'):
            stamps['h1'] = p
        elif not stamps['h2'] and p.style.name.startswith('Heading 2'):
            stamps['h2'] = p
        elif not stamps['h3'] and p.style.name.startswith('Heading 3'):
            stamps['h3'] = p
        elif not stamps['q'] and text_s and (p.style.name.startswith('021-') or p.style.name == 'List Paragraph'):
            stamps['q'] = p
        elif not stamps['normal'] and p.style.name == 'Normal' and text_s:
            stamps['normal'] = p
        elif not stamps['blank'] and p.style.name == 'Normal' and not text_s:
            stamps['blank'] = p
            
    for t in doc.tables:
        if not stamps['merge_table'] and len(t.columns) == 1 and len(t.rows) >= 2:
            stamps['merge_table'] = t
        elif len(t.columns) > 1:
            if not stamps['quiz_table'] and len(t.columns) == 2 and len(t.rows) > 0 and ("☐" in t.cell(0, 1).text or "主題句" in t.cell(0, 1).text):
                stamps['quiz_table'] = t
            elif not stamps['md_table']:
                stamps['md_table'] = t
            
    return stamps

def set_cant_split(table):
    for r in table.rows:
        trPr = r._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))

def clone_paragraph(stamps, doc, stamp_key, text=""):
    """
    利用 XML 深層複製技術克隆指定的段落屬性，保留段落格式並重寫 Runs，且解析 Markdown 語法
    """
    # 清理開頭的 Markdown 多餘符號（星號、減號、大於符號等），限定後面必須有空白，避免誤傷粗體雙星號 **
    cleaned_text = re.sub(r'^(?:[>\-\*—–]\s+)+', '', text)
    
    # 如果是提問段落，自動剃除手寫的數字題號（如 "1. " 或 "25. 3. "）-- V5.0 保留手寫題號，不交由 Word 清單自動編號以防黑點
    # if stamp_key == 'q':
    #     cleaned_text = re.sub(r'^\d+\s*[\.\-–—>]*\s*', '', cleaned_text)
        
    stamp_p = stamps.get(stamp_key)
    if stamp_p is None:
        p = doc.add_paragraph()
        format_paragraph(p, space_after_pt=0, line_spacing=1.0)
        if cleaned_text:
            add_formatted_text(p, cleaned_text, font_name="芫荽")
        return p
        
    p_element = copy.deepcopy(stamp_p._p)
    # 移除 List Paragraph / q 樣式的自動編號/黑點屬性 (w:numPr)，以配合手寫編號並消除黑點
    if stamp_key == 'q':
        pPr = p_element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
    
    # 處理 runs 文字重寫，清空舊的 runs
    runs = p_element.findall(qn('w:r'))
    if runs:
        for r in runs:
            p_element.remove(r)
            
    p = Paragraph(p_element, doc)
    append_element_safely(doc, p_element)
    
    # 取得範本段落的字型與顏色屬性
    font_name = "芫荽"
    font_size_pt = 11
    color_rgb = None
    
    if stamp_p.runs:
        r0 = stamp_p.runs[0]
        if r0.font.size:
            font_size_pt = r0.font.size.pt
        if r0.font.color and r0.font.color.rgb:
            color_rgb = r0.font.color.rgb
            
    # 只有 Heading 標題才預設加粗，其餘樣式（Normal, q, blank）一律預設非粗體，消除莫名粗體
    if stamp_key in ['h1', 'h2', 'h3']:
        default_bold = True
    else:
        default_bold = False
        
    # 設定行距：普通段落皆為 1.0 單行，且段落後間距為 0
    format_paragraph(p, space_after_pt=0, line_spacing=1.0)
    
    add_formatted_text(p, cleaned_text, default_bold=default_bold, font_name=font_name, font_size_pt=font_size_pt, color_rgb=color_rgb)
    return p

def clone_merge_table(stamps, doc, cached_data, cached_guide, cached_quiz):
    """
    克隆 1 欄多行的專屬合併表格 XML，並寫入新內容，100% 繼承原版表格框線與儲存格網底。
    """
    items = []
    if cached_data:
        items.append(('data', cached_data))
    if cached_guide:
        items.append(('guide', cached_guide))
    if cached_quiz:
        items.append(('quiz', cached_quiz))
        
    if not items:
        return
        
    stamp_t = stamps.get('merge_table')
    if stamp_t is None:
        table = doc.add_table(rows=len(items), cols=1, style='Table Grid')
        for idx, (item_type, lines) in enumerate(items):
            cell = table.cell(idx, 0)
            p = cell.paragraphs[0]
            format_paragraph(p, space_after_pt=0, line_spacing=1.0)
            run = p.add_run("\n".join(lines))
            apply_font(run, font_name="芫荽")
        set_cant_split(table)
        return
        
    tbl_element = copy.deepcopy(stamp_t._tbl)
    table = Table(tbl_element, doc)
    append_element_safely(doc, tbl_element)
    
    # 調整 rows
    while len(table.rows) > len(items):
        table._tbl.remove(table.rows[-1]._tr)
        
    # 填充內容
    for idx, (item_type, lines) in enumerate(items):
        cell = table.cell(idx, 0)
        
        # 100% 清空 cell 中除了第一個段落以外的所有段落與表格，防止範本文字殘留
        tc = cell._tc
        children = list(tc)
        first_p_found = False
        for child in children:
            if child.tag == qn('w:p'):
                if not first_p_found:
                    first_p_found = True
                else:
                    tc.remove(child)
            elif child.tag == qn('w:tbl'):
                tc.remove(child)
                
        p = cell.paragraphs[0]
        p.text = ""
        # 合併表格內部段落後間距 0，單行行距 1.0
        format_paragraph(p, space_before_pt=3, space_after_pt=0, line_spacing=1.0)
        
        if item_type == 'data':
            prefix = ""
        elif item_type == 'guide':
            prefix = "【教師引導】"
        else:
            prefix = "【學習評量】"
            
        full_text = "\n".join(lines).strip()
        # 移除開頭可能重複的標記前綴（含 Markdown 粗體 **），防止雙重前綴
        for _tag in ['【教師引導】', '【學習評量】']:
            _dup = r'^\*{0,2}' + re.escape(_tag) + r'\*{0,2}\s*'
            if re.match(_dup, full_text):
                full_text = re.sub(_dup, '', full_text, count=1).strip()
        
        first_line = True
        for line in (prefix + full_text if prefix and not full_text.startswith(prefix) else full_text).split("\n"):
            if not first_line:
                p = cell.add_paragraph()
                # 合併表格多行維持單行行距 1.0，段後間距 0
                format_paragraph(p, space_before_pt=2, space_after_pt=0, line_spacing=1.0)
            first_line = False
            
            # 清理開頭可能殘留的 markdown 標記（如 >, -, *, +）
            cleaned_line = re.sub(r'^(?:[>\-\*—–]\s+)+', '', line)
            
            # 答案與解析過濾（防呆）
            if "正確答案" in cleaned_line or "解析" in cleaned_line or "選題理由" in cleaned_line or "出處" in cleaned_line:
                continue
                
            if cleaned_line.strip().startswith(("【學習評量】", "【教師引導】", "【資料")):
                match = re.match(r'^(【.*?】)(.*)', cleaned_line)
                if match:
                    r_bold = p.add_run(match.group(1))
                    apply_font(r_bold, font_name="芫荽", font_size_pt=10.5, bold=True)
                    add_formatted_text(p, match.group(2), default_bold=False, font_size_pt=10.5)
                else:
                    add_formatted_text(p, cleaned_line, default_bold=False, font_size_pt=10.5)
            else:
                add_formatted_text(p, cleaned_line, default_bold=False, font_size_pt=10.5)
                
    set_cant_split(table)
    clone_paragraph(stamps, doc, 'blank')

def clone_markdown_table(stamps, doc, table_rows):
    """
    克隆普通多欄分析表格 XML，並動態填寫，100% 繼承分析表格框線與儲存格樣式。
    手寫填寫表格的資料列設定為 1.5 倍行高。
    """
    if not table_rows:
        return
        
    parsed_rows = []
    for row in table_rows:
        parts = [p.strip() for p in row.strip().strip('|').split('|')]
        parsed_rows.append(parts)
        
    if not parsed_rows:
        return
        
    num_rows = len(parsed_rows)
    num_cols = max(len(r) for r in parsed_rows)
    
    stamp_t = stamps.get('md_table')
    if stamp_t is None:
        table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
        for r_idx, row_data in enumerate(parsed_rows):
            row = table.rows[r_idx]
            for c_idx, val in enumerate(row_data):
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    p = cell.paragraphs[0]
                    line_spacing = 1.0 if r_idx == 0 else 1.5
                    format_paragraph(p, space_after_pt=0, line_spacing=line_spacing)
                    add_formatted_text(p, val, font_name="芫荽")
        set_cant_split(table)
        return
        
    tbl_element = copy.deepcopy(stamp_t._tbl)
    table = Table(tbl_element, doc)
    append_element_safely(doc, tbl_element)
    
    # 調整 rows
    while len(table.rows) < num_rows:
        new_tr = copy.deepcopy(table.rows[-1]._tr)
        table._tbl.append(new_tr)
    while len(table.rows) > num_rows:
        table._tbl.remove(table.rows[-1]._tr)
        
    # 填充內容
    for r_idx, row_data in enumerate(parsed_rows):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx, val in enumerate(row_data):
            if c_idx < len(row.cells):
                cell = row.cells[c_idx]
                
                # 100% 清空 cell 中除了第一個段落以外的所有段落與表格，防止範本文字殘留
                tc = cell._tc
                children = list(tc)
                first_p_found = False
                for child in children:
                    if child.tag == qn('w:p'):
                        if not first_p_found:
                            first_p_found = True
                        else:
                            tc.remove(child)
                    elif child.tag == qn('w:tbl'):
                        tc.remove(child)
                        
                p = cell.paragraphs[0]
                p.text = ""
                # 設定行距：表頭（Header）為單行 1.0 行距；手寫填寫列為 1.5 倍行高，且段落後間距皆為 0
                line_spacing = 1.0 if is_header else 1.5
                format_paragraph(p, space_before_pt=3, space_after_pt=0, line_spacing=line_spacing)
                
                cleaned_val = val.strip()
                add_formatted_text(p, cleaned_val, default_bold=is_header, font_size_pt=10.5)
                
    set_cant_split(table)
    clone_paragraph(stamps, doc, 'blank')

def clone_quiz_table(stamps, doc, quizzes):
    """
    克隆後測 2 欄專屬表格 XML，將問題填寫於左欄，檢核清單維持在右欄。
    """
    if not quizzes:
        return
        
    stamp_t = stamps.get('quiz_table')
    if stamp_t is None:
        for q in quizzes:
            clone_paragraph(stamps, doc, 'q', text=q)
            for _ in range(3):
                clone_paragraph(stamps, doc, 'blank')
        return
        
    num_rows = len(quizzes)
    tbl_element = copy.deepcopy(stamp_t._tbl)
    table = Table(tbl_element, doc)
    append_element_safely(doc, tbl_element)
    
    while len(table.rows) < num_rows:
        new_tr = copy.deepcopy(table.rows[-1]._tr)
        table._tbl.append(new_tr)
    while len(table.rows) > num_rows:
        table._tbl.remove(table.rows[-1]._tr)
        
    for idx, q_text in enumerate(quizzes):
        row = table.rows[idx]
        
        # 1. 處理左欄：題目
        cell_q = row.cells[0]
        tc_q = cell_q._tc
        for child in list(tc_q):
            if child.tag == qn('w:p'):
                tc_q.remove(child)
            elif child.tag == qn('w:tbl'):
                tc_q.remove(child)
                
        p_q = Paragraph(OxmlElement('w:p'), doc)
        tc_q.append(p_q._p)
        format_paragraph(p_q, space_before_pt=4, space_after_pt=4, line_spacing=1.0)
        
        # 標準化為中文題號 (一、二、三...)
        chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        cleaned_q = re.sub(r'^\d+\s*[\.\-–—>]*\s*', '', q_text).strip()
        if idx < len(chinese_nums):
            prefix = f"{chinese_nums[idx]}、"
        else:
            prefix = f"{idx+1}、"
            
        add_formatted_text(p_q, prefix + cleaned_q, default_bold=False, font_size_pt=10.5)
        
        # 2. 處理右欄：檢核清單
        cell_a = row.cells[1]
        for p in cell_a.paragraphs:
            format_paragraph(p, space_before_pt=2, space_after_pt=2, line_spacing=1.0)
            for r in p.runs:
                apply_font(r, font_name="芫荽", font_size_pt=9.5)
                
    set_cant_split(table)
    clone_paragraph(stamps, doc, 'blank')

def update_paragraph_header_text(p, new_title):
    """
    精確替換頁首段落中的課名，保留定位符 \t 與右側的姓名座號
    """
    text = p.text
    # 尋找是否含有舊課名關鍵字或符合 "第 X 課" 的特徵
    for keyword in ["成功是失敗之母", "我所知道的康橋", "康橋", "項鍊", "陋室銘", "街亭", "空城計", "第"]:
        if keyword in text:
            # 尋找第一個定位點 \t 在哪個 run
            tab_index = -1
            for r_idx, run in enumerate(p.runs):
                if '\t' in run.text:
                    tab_index = r_idx
                    break
            
            if tab_index != -1:
                # 找到了定位符。我們把 0 到 tab_index - 1 的 runs 的文字全部設為 ""
                # 並把 new_title 寫入 run 0
                for r_idx in range(tab_index):
                    p.runs[r_idx].text = ""
                p.runs[0].text = new_title
                apply_font(p.runs[0], font_name="芫荽")
            else:
                # 如果沒有定位點，直接替換段落 text 或者是將第一個 run 填入，其他設為 ""
                if len(p.runs) > 0:
                    for r_idx in range(1, len(p.runs)):
                        p.runs[r_idx].text = ""
                    p.runs[0].text = new_title
                    apply_font(p.runs[0], font_name="芫荽")
            break

def update_document_headers(doc, new_title):
    """
    遍歷文件中所有 sections 的所有頁首（普通、首頁、偶數頁）
    """
    for section in doc.sections:
        header_objs = [section.header, section.first_page_header, section.even_page_header]
        for header in header_objs:
            if header is None:
                continue
            for p in header.paragraphs:
                update_paragraph_header_text(p, new_title)

def parse_md_to_docx(md_path, template_path, output_path):
    # 開啟範本
    doc_template = docx.Document(template_path)
    # 提取印章
    stamps = extract_stamps(doc_template)
    
    # 新建空白文件，承接樣式
    doc_new = docx.Document(template_path)
    clean_document_content(doc_new)
    
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    # 提取課名
    lesson_title = "自學3 空城計"
    for line in lines:
        line_str = line.strip()
        if line_str.startswith('# '):
            raw_title = line_str.replace('#', '').strip()
            for suffix in ["學思達教學講義", "學思達講義", "教學講義", "講義"]:
                raw_title = raw_title.replace(suffix, "")
            lesson_title = raw_title.strip()
            break
            
    # 先寫入大標題 (Heading 1) 且不加手寫「壹、」前綴，防 Word 自動編號重疊
    clone_paragraph(stamps, doc_new, 'h1', text=f"{lesson_title} 學思達教學講義")
    
    cached_data = []
    cached_guide = []
    cached_quiz = []
    cached_table_rows = []
    cached_aftertest_quizzes = []
    
    current_section = None
    
    def flush_all_caches():
        nonlocal cached_table_rows
        if cached_data or cached_guide or cached_quiz:
            clone_merge_table(stamps, doc_new, cached_data, cached_guide, cached_quiz)
            cached_data.clear()
            cached_guide.clear()
            cached_quiz.clear()
        if cached_table_rows:
            clone_markdown_table(stamps, doc_new, cached_table_rows)
            cached_table_rows.clear()
        if cached_aftertest_quizzes:
            clone_quiz_table(stamps, doc_new, cached_aftertest_quizzes)
            cached_aftertest_quizzes.clear()
            
    start_output = False
    
    for line in lines:
        line_str = line.strip()
        
        # 答案、解析與選題理由自動過濾
        if "正確答案" in line_str or "【解析】" in line_str or "解析：" in line_str or "選題理由" in line_str or "【出處】" in line_str:
            continue
            
        # 尋找「## 壹、暖身題」作為開始輸出的標記，在此之前的所有教學設計藍圖內容皆略過
        if not start_output:
            if line_str.startswith('##') and '暖身題' in line_str:
                start_output = True
                # 找到暖身題了，開始正常寫入
            else:
                continue
                
        # === 過濾規則 ===
        # 跳過 Markdown 水平分隔線及自訂分隔線 (如 ---A方案---)
        if re.match(r'^-{3,}', line_str):
            continue
        
        # 跳過純引用標記行（如獨立的 > 符號）
        if re.match(r'^>+\s*$', line_str):
            continue
        
        # 跳過教師專用段落（不應出現在學生講義上）
        _TEACHER_ONLY = [
            '(Tier ', '（Tier ',
            '【設計理念】', '【與課文連結點】',
            '【活動/提問名稱】', '【活動/提問內容】',
            '【建議討論鷹架】', '【評分標準提示】', '【出處】',
        ]
        _cleaned_check = re.sub(r'^[\*>#\s]+', '', line_str)
        if any(_cleaned_check.startswith(pfx) for pfx in _TEACHER_ONLY):
            # 設計理念/連結點/鷹架/評分標準 後方有教師備註文字，需一併跳過
            _SKIP_FOLLOWING = ['【設計理念】', '【與課文連結點】', '【建議討論鷹架】', '【評分標準提示】']
            if any(_cleaned_check.startswith(s) for s in _SKIP_FOLLOWING):
                current_section = 'skip_teacher'
            continue
        
        # 一級標題
        h1_match = re.match(r'^#\s+(.*)', line_str)
        if h1_match:
            flush_all_caches()
            current_section = 'normal'
            clone_paragraph(stamps, doc_new, 'h1', text=h1_match.group(1))
            continue
            
        # 二級標題
        h2_match = re.match(r'^##\s+(.*)', line_str)
        if h2_match:
            flush_all_caches()
            title_raw = h2_match.group(1).strip('* ')
            title_text = re.sub(r'^[壹貳參肆伍陸柒捌玖拾]+、\s*', '', title_raw)
            
            # 標準化 Heading 1 大標題文字，100% 對齊 3W 範例檔格式
            if "暖身題" in title_text:
                title_text = "暖身題"
                current_section = 'normal'
            elif "課文理解" in title_text:
                title_text = "課文理解"
                current_section = 'normal'
            elif "統整探究" in title_text:
                title_text = "統整探究"
                current_section = 'normal'
            elif "課後延伸" in title_text:
                title_text = "課後延伸"
                current_section = 'normal'
            elif "後測" in title_text:
                title_text = "後測"
                current_section = 'aftertest'
            else:
                current_section = 'normal'
                
            clone_paragraph(stamps, doc_new, 'h1', text=title_text)
            continue
            
        # 三級標題
        h3_match = re.match(r'^###\s+(.*)', line_str)
        if h3_match:
            flush_all_caches()
            current_section = 'normal'
            title_text = h3_match.group(1).strip('* ')
            clone_paragraph(stamps, doc_new, 'h2', text=title_text)
            continue
            
        # What / How / Why 標題偵測與標準化
        whw_match = re.match(r'^[\(【]?(what|how|why)(?:\s*[:：]\s*(?:文意理解|寫作手法|作者意圖))?[\)】]?$', line_str, re.IGNORECASE)
        if whw_match:
            flush_all_caches()
            current_section = 'normal'
            whw_type = whw_match.group(1).lower()
            if whw_type == 'what':
                label = '【What：文意理解】'
            elif whw_type == 'how':
                label = '【How：寫作手法】'
            else:
                label = '【Why：作者意圖】'
            
            p = clone_paragraph(stamps, doc_new, 'normal', text=label)
            p.runs[0].bold = True
            format_paragraph(p, space_before_pt=6, space_after_pt=0, line_spacing=1.0)
            continue
            
        # 合併表格標籤
        h4_match = re.match(r'^####\s+(.*)', line_str)
        if h4_match:
            tag = h4_match.group(1).strip()
            
            # 判斷是否為合併表格的三種標籤
            is_merge_tag = False
            target_type = None
            if "資料" in tag or "補充" in tag:
                is_merge_tag = True
                target_type = 'data'
            elif "引導" in tag:
                is_merge_tag = True
                target_type = 'guide'
            elif "評量" in tag:
                is_merge_tag = True
                target_type = 'quiz'
                
            if is_merge_tag:
                # 根據順序判定是否需要 flush，若是全新合併表格才 flush
                need_flush = False
                if target_type == 'data':
                    if cached_data or cached_guide or cached_quiz:
                        need_flush = True
                elif target_type == 'guide':
                    if cached_guide or cached_quiz:
                        need_flush = True
                elif target_type == 'quiz':
                    if cached_quiz:
                        need_flush = True
                        
                if need_flush:
                    flush_all_caches()
                
                current_section = target_type
            else:
                # 不是合併表格標籤，是一般四級標題，必須先 flush 前面的快取
                flush_all_caches()
                current_section = 'normal'
                # 跳過子結構標題（如「（一）提問 (層次性問題)」）
                cleaned_tag = tag.strip('* ').strip()
                if re.match(r'^[（\(][一二三四五六七八九十]+[）\)]', cleaned_tag):
                    continue
                # 一般小標題，直接寫入為普通內文
                clone_paragraph(stamps, doc_new, 'normal', text=tag)
            continue
            
        # 普通表格行
        if line_str.startswith('|'):
            # 修復分割線正則表達式中的減號 Bug，對減號和豎線進行正確跳脫
            if re.match(r'^\|[\s:\-\|]+$', line_str):
                continue
            cached_table_rows.append(line_str)
            continue
        elif cached_table_rows and not line_str.startswith('|'):
            clone_markdown_table(stamps, doc_new, cached_table_rows)
            cached_table_rows.clear()
            
        if current_section == 'data':
            if line_str:
                cached_data.append(line_str)
            continue
        elif current_section == 'guide':
            if line_str:
                cached_guide.append(line_str)
            continue
        elif current_section == 'quiz':
            if line_str:
                cached_quiz.append(line_str)
            continue
        elif current_section == 'skip_teacher':
            continue
            
        if not line_str:
            continue
            
        # 提問清單
        quiz_match = re.match(r'^(\d+\..*)', line_str)
        if quiz_match:
            if current_section == 'aftertest':
                cached_aftertest_quizzes.append(line_str)
            else:
                # 去除數字點編號，判斷是否為標題（避免 1. 【主旨探究】 等大題子標題誤判為普通提問而留白）
                rest = quiz_match.group(1).split('.', 1)[1].strip()
                rest_clean = rest.lstrip('*').strip()
                if rest_clean.startswith('【'):
                    # 這是一般標題，直接作為 normal 寫入，且不留白
                    clone_paragraph(stamps, doc_new, 'normal', text=line_str)
                else:
                    p = clone_paragraph(stamps, doc_new, 'q', text=line_str)
                    p.paragraph_format.keep_with_next = True
                    
                    # 插入 3 個空白段落，前 2 個空白段落也設定 keep_with_next = True
                    for step in range(3):
                        p_empty = clone_paragraph(stamps, doc_new, 'blank')
                        keep = (step < 2)
                        p_empty.paragraph_format.keep_with_next = keep
        else:
            # 一般內文
            clone_paragraph(stamps, doc_new, 'normal', text=line_str)
            
    # 處理剩餘快取
    flush_all_caches()
    
    # 動態更新頁首
    update_document_headers(doc_new, lesson_title)
    
    # 儲存
    doc_new.save(output_path)
    print(f"Successfully generated Word file using XML Stamps: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert 学思达 handout Markdown to Word (.docx) using Style Stamps.")
    parser.add_argument("--md", required=True, help="Path to input Markdown file")
    parser.add_argument("--out", required=True, help="Path to output Word file")
    parser.add_argument("--template", default=r"C:\Users\sshor\.gemini\config\skills\shareclass-handout\assets\template.docx", help="Path to template Docx file")
    args = parser.parse_args()
    
    parse_md_to_docx(args.md, args.template, args.out)
